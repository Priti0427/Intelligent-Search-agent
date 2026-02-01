"""
Document Loader.

This module handles loading documents from various file formats
including PDF, DOCX, TXT, and Markdown files.
"""

import logging
from pathlib import Path
from typing import List, Optional, Union

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)

logger = logging.getLogger(__name__)


class Document:
    """A loaded document with content and metadata."""
    
    def __init__(self, content: str, metadata: dict):
        self.content = content
        self.metadata = metadata
    
    def __repr__(self):
        return f"Document(source={self.metadata.get('source', 'unknown')}, len={len(self.content)})"


class DocumentLoader:
    """
    Multi-format document loader.
    
    Supports:
    - PDF files (.pdf)
    - Text files (.txt)
    - Markdown files (.md)
    - Word documents (.docx) - requires python-docx
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx"}
    
    def __init__(self, base_path: Optional[str] = None):
        """
        Initialize the document loader.
        
        Args:
            base_path: Base directory for relative paths
        """
        self.base_path = Path(base_path) if base_path else Path("./data/documents")
    
    def load_file(self, file_path: Union[str, Path]) -> List[Document]:
        """
        Load a single file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document objects (may be multiple for multi-page PDFs)
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {path}")
            return []
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            logger.warning(f"Unsupported file type: {extension}")
            return []
        
        logger.info(f"Loading file: {path}")
        
        try:
            if extension == ".pdf":
                return self._load_pdf(path)
            elif extension in {".txt"}:
                return self._load_text(path)
            elif extension in {".md", ".markdown"}:
                return self._load_markdown(path)
            elif extension == ".docx":
                return self._load_docx(path)
            else:
                return []
                
        except Exception as e:
            logger.error(f"Failed to load {path}: {e}")
            return []
    
    def _load_pdf(self, path: Path) -> List[Document]:
        """Load a PDF file."""
        loader = PyPDFLoader(str(path))
        pages = loader.load()
        
        documents = []
        for i, page in enumerate(pages):
            documents.append(
                Document(
                    content=page.page_content,
                    metadata={
                        "source": str(path),
                        "title": path.stem,
                        "page": i + 1,
                        "total_pages": len(pages),
                        "file_type": "pdf",
                    },
                )
            )
        
        return documents
    
    def _load_text(self, path: Path) -> List[Document]:
        """Load a text file."""
        loader = TextLoader(str(path))
        docs = loader.load()
        
        return [
            Document(
                content=doc.page_content,
                metadata={
                    "source": str(path),
                    "title": path.stem,
                    "file_type": "txt",
                },
            )
            for doc in docs
        ]
    
    def _load_markdown(self, path: Path) -> List[Document]:
        """Load a Markdown file."""
        try:
            loader = UnstructuredMarkdownLoader(str(path))
            docs = loader.load()
            
            return [
                Document(
                    content=doc.page_content,
                    metadata={
                        "source": str(path),
                        "title": path.stem,
                        "file_type": "markdown",
                    },
                )
                for doc in docs
            ]
        except Exception:
            # Fallback to simple text loading
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            
            return [
                Document(
                    content=content,
                    metadata={
                        "source": str(path),
                        "title": path.stem,
                        "file_type": "markdown",
                    },
                )
            ]
    
    def _load_docx(self, path: Path) -> List[Document]:
        """Load a Word document."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(path))
            content = "\n\n".join([para.text for para in doc.paragraphs if para.text])
            
            return [
                Document(
                    content=content,
                    metadata={
                        "source": str(path),
                        "title": path.stem,
                        "file_type": "docx",
                    },
                )
            ]
        except ImportError:
            logger.error("python-docx not installed. Cannot load .docx files.")
            return []
    
    def load_directory(
        self,
        directory: Optional[Union[str, Path]] = None,
        recursive: bool = True,
    ) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Directory path (uses base_path if not provided)
            recursive: Whether to search subdirectories
            
        Returns:
            List of all loaded documents
        """
        dir_path = Path(directory) if directory else self.base_path
        
        if not dir_path.exists():
            logger.warning(f"Directory not found: {dir_path}")
            return []
        
        logger.info(f"Loading documents from: {dir_path}")
        
        documents = []
        
        # Get all files
        if recursive:
            files = list(dir_path.rglob("*"))
        else:
            files = list(dir_path.glob("*"))
        
        # Filter to supported extensions
        files = [f for f in files if f.suffix.lower() in self.SUPPORTED_EXTENSIONS]
        
        logger.info(f"Found {len(files)} supported files")
        
        for file_path in files:
            docs = self.load_file(file_path)
            documents.extend(docs)
        
        logger.info(f"Loaded {len(documents)} documents total")
        return documents


def load_documents(
    path: Union[str, Path],
    recursive: bool = True,
) -> List[Document]:
    """
    Convenience function to load documents from a path.
    
    Args:
        path: File or directory path
        recursive: Whether to search subdirectories (for directories)
        
    Returns:
        List of loaded documents
    """
    loader = DocumentLoader()
    path = Path(path)
    
    if path.is_file():
        return loader.load_file(path)
    elif path.is_dir():
        return loader.load_directory(path, recursive=recursive)
    else:
        logger.error(f"Path not found: {path}")
        return []
